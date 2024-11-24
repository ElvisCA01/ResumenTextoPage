import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from collections import defaultdict
import string
import requests
from bs4 import BeautifulSoup
import re
import fitz  # PyMuPDF
import torch
from transformers import pipeline
import docx
import odf.text
import odf.teletype
from odf.opendocument import load
import os
import chardet

nltk.download('punkt_tab')
nltk.download('punkt')
nltk.download('stopwords')

# Descargar recursos de NLTK si es necesario
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Verificar si hay GPU disponible
device = 0 if torch.cuda.is_available() else -1

# Inicializar el pipeline de summarization
try:
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=device)
except Exception as e:
    print(f"Error al cargar el modelo: {str(e)}")
    summarizer = None

def generate_summary(text, max_length=100, min_length=50):
    """
    Genera un resumen utilizando el pipeline de transformers.
    """
    try:
        if summarizer is None:
            return "Error: El modelo no se pudo cargar correctamente."

        # Dividir el texto en chunks si es muy largo
  
        chunks = dividir_en_chunks(text)
        
        summaries = []
        for chunk in chunks:
            if len(chunk.strip()) > 100:  # Solo procesar chunks con suficiente contenido
                summary = summarizer(chunk, 
                                   max_length=max_length, 
                                   min_length=min_length, 
                                   do_sample=False)
                summaries.append(summary[0]['summary_text'])
        
        return ' '.join(summaries)
    except Exception as e:
        return f"Error al generar el resumen: {str(e)}"

# Funcion mejorada para la division de chunks
def dividir_en_chunks(texto, max_chunk_length=1024):
    """
    Divide el texto en chunks sin cortar oraciones.
    """
    chunks = []
    actual = ""
    for oracion in sent_tokenize(texto):
        if len(actual) + len(oracion) <= max_chunk_length:
            actual += " " + oracion
        else:
            chunks.append(actual.strip())
            actual = oracion
    if actual:
        chunks.append(actual.strip())
    return chunks

def ajustar_resumen_por_tamano(resumen, tamano):
    """
    Ajusta el tamaño del resumen según la opción seleccionada.
    """
    if not resumen:
        return ""
        
    oraciones = sent_tokenize(resumen)
    
    if tamano == 'corto':
        num_oraciones = min(2, len(oraciones))
    elif tamano == 'medio':
        num_oraciones = min(5, len(oraciones))
    else:  # largo
        num_oraciones = min(10, len(oraciones))
    
    return ' '.join(oraciones[:num_oraciones])


def limpiar_texto(texto):
    """
    Limpia el texto eliminando referencias, números, múltiples espacios, etc.
    """
    if not texto:
        return ""
    texto = re.sub(r'\[\d+\]', '', texto)  # Eliminar referencias
    texto = re.sub(r'[\r\n]+', ' ', texto)  # Reemplazar saltos de línea por espacios
    texto = re.sub(r'http\S+', '', texto)   # Quitar URLs
    texto = re.sub(r'\s+', ' ', texto) #Eliminar espacios multiples
    return texto.strip()


def resumen_nltk(texto, tamano):
    """
    Genera un resumen simple utilizando NLTK y ajusta el tamaño según la opción seleccionada.
    """
    texto_limpio = limpiar_texto(texto)
    sentences = sent_tokenize(texto_limpio)
    stop_words = set(stopwords.words('spanish'))
    words = [word for word in word_tokenize(texto_limpio.lower())
             if word.isalnum() and word not in stop_words]
    
    frequency = defaultdict(int)
    for word in words:
        frequency[word] += 1

    sentence_score = defaultdict(int)
    for sentence in sentences:
        for word in word_tokenize(sentence.lower()):
            if word in frequency:
                sentence_score[sentence] += frequency[word]

    top_sentences = sorted(sentence_score, key=sentence_score.get, reverse=True)[:int(len(sentences) * 0.3)]
    resumen = ' '.join(sorted(top_sentences, key=lambda x: sentences.index(x)))

    # Ajustamos el tamaño del resumen
    return ajustar_resumen_por_tamano(resumen, tamano)


def detectar_codificacion(contenido):
    """
    Detecta la codificación del archivo de texto.
    """
    resultado = chardet.detect(contenido)
    return resultado['encoding']

def obtener_texto_de_archivo(archivo):
    """
    Extrae texto de diferentes tipos de archivos.
    Soporta: .txt, .pdf, .doc, .docx, .odt, .rtf
    """
    try:
        nombre_archivo = archivo.name.lower()
        extension = os.path.splitext(nombre_archivo)[1]
        
        # Asegurarse de que estamos al inicio del archivo
        archivo.seek(0)
        
        if extension == '.pdf':
            return obtener_texto_de_pdf(archivo)
            
        elif extension == '.docx':
            return obtener_texto_de_docx(archivo)
            
        elif extension == '.odt':
            return obtener_texto_de_odt(archivo)
            
        elif extension == '.txt':
            return obtener_texto_de_txt(archivo)
            
        elif extension == '.rtf':
            return obtener_texto_de_rtf(archivo)
            
        else:
            # Intentar leer como texto plano si no coincide con ningún formato conocido
            try:
                return obtener_texto_de_txt(archivo)
            except:
                raise Exception(f"Formato de archivo no soportado: {extension}")
                
    except Exception as e:
        return f"Error al procesar el archivo: {str(e)}"


def obtener_texto_de_docx(archivo):
    """
    Extrae texto de un archivo DOCX.
    """
    doc = docx.Document(archivo)
    texto_completo = []
    
    for parrafo in doc.paragraphs:
        texto_completo.append(parrafo.text)
        
    # Extraer texto de las tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                texto_completo.append(celda.text)
    
    return '\n'.join(texto_completo)

def obtener_texto_de_odt(archivo):
    """
    Extrae texto de un archivo ODT.
    """
    textdoc = load(archivo)
    allparas = textdoc.getElementsByType(odf.text.P)
    return '\n'.join([odf.teletype.extractText(para) for para in allparas])

def obtener_texto_de_txt(archivo):
    """
    Extrae texto de un archivo TXT.
    """
    contenido = archivo.read()
    if isinstance(contenido, str):
        return contenido.strip()
    
    # Si el contenido es bytes, detectar la codificación
    codificacion = detectar_codificacion(contenido)
    return contenido.decode(codificacion or 'utf-8', errors='replace').strip()

def obtener_texto_de_rtf(archivo):
    """
    Extrae texto de un archivo RTF.
    """
    try:
        import striprtf.striprtf as striprtf
        contenido = archivo.read()
        if isinstance(contenido, bytes):
            contenido = contenido.decode('utf-8', errors='replace')
        return striprtf.rtf_to_text(contenido).strip()
    except ImportError:
        return "Error: Se requiere la biblioteca 'striprtf' para procesar archivos RTF"

# Actualizar la función de validación de archivos
def es_archivo_valido(nombre_archivo):
    """
    Verifica si el archivo tiene una extensión válida.
    """
    extensiones_permitidas = {
        '.txt', '.pdf', '.doc', '.docx', 
        '.odt', '.rtf'
    }
    extension = os.path.splitext(nombre_archivo.lower())[1]
    return extension in extensiones_permitidas

def obtener_texto_de_pdf(archivo):
    """
    Extrae texto de un archivo PDF utilizando PyMuPDF.
    """
    try:
        texto = ""
        # Asegurarse de que estamos al inicio del archivo
        archivo.seek(0)
        
        # Leer el contenido del archivo
        pdf_bytes = archivo.read()
        
        # Abrir el PDF desde los bytes
        with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf_document:
            # Extraer texto de cada página
            for pagina in pdf_document:
                texto += pagina.get_text()
        
        return texto.strip()
    except Exception as e:
        return f"Error al procesar el PDF: {str(e)}"

def obtener_texto_de_url(url, max_length=2000):
    """
    Obtiene texto de una URL extrayendo contenido de etiquetas <p> y limita el tamaño.
    Se agrega un parámetro max_length para limitar la cantidad de texto que se usa para el resumen.
    """
    response = requests.get(url)
    response.raise_for_status()
    # Intentar obtener la codificación desde el encabezado de la respuesta
    response.encoding = 'utf-8'
    soup = BeautifulSoup(response.text, 'html.parser')
    texto = ' '.join([para.get_text() for para in soup.find_all('p')])
    
    # Limitar la cantidad de texto que se toma (max_length caracteres)
    texto_limited = texto[:max_length]  # Limita a max_length caracteres
    return texto_limited

def procesar_entrada(input_type, contenido):
    """
    Procesa entrada desde texto, archivo o URL.
    """
    if input_type == 'texto':
        return contenido
    elif input_type == 'archivo':
        return obtener_texto_de_pdf(contenido)  # Procesar PDF
    elif input_type == 'url':
        return obtener_texto_de_url(contenido)  # Procesar URL
